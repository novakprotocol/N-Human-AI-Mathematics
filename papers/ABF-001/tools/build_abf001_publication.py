#!/usr/bin/env python3
"""Build reviewable ABF-001 PDF/DOCX assets from committed Markdown.

The Markdown source remains controlling. These are convenience renderings for
public technical review. The builder is deterministic at the text/layout level;
ZIP and PDF metadata are normalized by the release workflow where practical.
"""
from __future__ import annotations

import argparse
import html
import re
import subprocess
from pathlib import Path

import mistune
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from weasyprint import HTML

CSS = r"""
@page { size: Letter; margin: 0.72in 0.72in 0.72in 0.72in; @bottom-center { content: counter(page); font-size: 8pt; color: #555; } }
body { font-family: 'DejaVu Sans', Arial, sans-serif; font-size: 9.2pt; line-height: 1.38; color: #17202a; }
h1 { font-size: 20pt; margin: 0 0 8pt; line-height: 1.12; color: #102a43; }
h2 { font-size: 14pt; margin: 17pt 0 6pt; border-bottom: 1px solid #b8c7d9; padding-bottom: 2pt; color: #153e75; }
h3 { font-size: 11pt; margin: 12pt 0 4pt; color: #1f4e79; }
p { margin: 4pt 0; }
table { border-collapse: collapse; width: 100%; margin: 7pt 0 9pt; font-size: 8pt; break-inside: avoid; }
th, td { border: 0.6pt solid #aeb8c2; padding: 3.5pt; vertical-align: top; }
th { background: #e8eef5; font-weight: 700; }
pre { white-space: pre-wrap; font-family: 'DejaVu Sans Mono', monospace; background: #f4f6f8; border-left: 3pt solid #5b7c99; padding: 5pt; font-size: 7.7pt; }
code { font-family: 'DejaVu Sans Mono', monospace; font-size: 0.92em; }
blockquote { margin: 7pt 18pt; padding-left: 9pt; border-left: 3pt solid #7c9abb; color: #34495e; }
ul, ol { margin-top: 3pt; margin-bottom: 4pt; }
li { margin: 1.5pt 0; }
.math { font-family: 'DejaVu Serif', serif; }
"""


def clean_markdown(text: str, anonymous: bool) -> str:
    if anonymous:
        text = re.sub(r"(?m)^\*\*Author:\*\*.*\n", "", text)
        text = text.replace("Matthew S. Novak", "Anonymous author")
    return text


def markdown_html(text: str, title: str) -> str:
    renderer = mistune.create_markdown(plugins=["table", "strikethrough", "task_lists", "url"])
    body = renderer(text)
    return f"""<!doctype html><html lang='en'><head><meta charset='utf-8'><title>{html.escape(title)}</title><style>{CSS}</style></head><body>{body}</body></html>"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar"); fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar"); fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def build_docx(md_text: str, output: Path, anonymous: bool) -> None:
    text = clean_markdown(md_text, anonymous)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65); sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"; styles["Normal"].font.size = Pt(9)
    for name, size in (("Title", 20), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10)):
        styles[name].font.name = "Aptos Display"; styles[name].font.size = Pt(size)
    add_page_number(sec.footer.paragraphs[0])

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1; continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title"); p.add_run(line[2:].strip())
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
        elif line.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i]); i += 1
            p = doc.add_paragraph()
            p.style = styles["Normal"]
            r = p.add_run("\n".join(block)); r.font.name = "Aptos Mono"; r.font.size = Pt(8)
            p.paragraph_format.left_indent = Inches(0.18)
        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i+1]):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i += 1
            header = rows[0]
            data = rows[2:]
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for j, value in enumerate(header):
                table.rows[0].cells[j].text = re.sub(r"[*`]", "", value)
                set_cell_shading(table.rows[0].cells[j], "DCE6F1")
            for row in data:
                cells = table.add_row().cells
                for j, value in enumerate(row[:len(cells)]):
                    cells[j].text = re.sub(r"[*`]", "", value)
            i -= 1
        elif re.match(r"^[-*] ", line):
            doc.add_paragraph(re.sub(r"^[-*] ", "", line), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:]); p.paragraph_format.left_indent = Inches(0.25)
        else:
            p = doc.add_paragraph()
            # Lightweight bold/code cleanup while retaining readable mathematical source.
            parts = re.split(r"(\*\*.*?\*\*|`.*?`)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1]); run.font.name = "Aptos Mono"; run.font.size = Pt(8)
                else:
                    p.add_run(part)
        i += 1

    core = doc.core_properties
    core.title = "ABF-001"; core.subject = "Candidate public technical review"
    core.author = "Anonymous" if anonymous else "Matthew S. Novak"
    core.keywords = "Boolean functions; affine restrictions; Reed-Muller; incidence geometry"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_pair(md_path: Path, output_dir: Path, stem: str, anonymous: bool = False) -> None:
    md = md_path.read_text(encoding="utf-8")
    clean = clean_markdown(md, anonymous)
    html_text = markdown_html(clean, stem)
    HTML(string=html_text, base_url=str(md_path.parent)).write_pdf(output_dir / f"{stem}.pdf")
    build_docx(md, output_dir / f"{stem}.docx", anonymous)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, default=Path(__file__).resolve().parents[1])
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    root = args.root.resolve(); out = args.output.resolve(); out.mkdir(parents=True, exist_ok=True)
    build_pair(root / "manuscript/ABF-001_MANUSCRIPT.md", out, "ABF-001_Manuscript")
    build_pair(root / "manuscript/ABF-001_MANUSCRIPT.md", out, "ABF-001_Anonymous_Manuscript", True)
    build_pair(root / "review/ABF-001_REVIEWER_GUIDE.md", out, "ABF-001_Reviewer_Guide")
    build_pair(root / "review/ABF-001_EXECUTIVE_SUMMARY.md", out, "ABF-001_Executive_Summary")
    print(f"built {len(list(out.iterdir()))} publication assets in {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
